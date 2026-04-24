"""DOCX 正文/页眉/页脚/文本框/表格提取 (C5 parser-pipeline)

python-docx 读段落 + header + footer + textbox + table rows。
返回结构化列表,上游 __init__.extract_content 写 document_texts。
"""

from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import TYPE_CHECKING

from lxml import etree

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument  # type: ignore

logger = logging.getLogger(__name__)

# parser-accuracy-fixes P2-8:预编译 lxml XPath,绕过 python-docx BaseOxmlElement.xpath(namespaces=...)
# 新版 python-docx 的 BaseOxmlElement.xpath() 已废 `namespaces` kwarg(抛 TypeError)
# lxml.etree.XPath 直接传 namespaces 编译,对 BaseOxmlElement 底层 lxml Element 可用
_TEXTBOX_XPATH = etree.XPath(
    ".//w:txbxContent//w:t",
    namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
)


@dataclass(frozen=True)
class TextBlock:
    paragraph_index: int
    text: str
    # body | header | footer | textbox | table_row
    location: str


@dataclass(frozen=True)
class DocxExtractResult:
    blocks: list[TextBlock]
    # 提取阶段的 warning(非致命),上游可 log 但不写 parse_error
    warnings: list[str]


def extract_docx(file_path: str | Path) -> DocxExtractResult:
    """从 docx 抽文本。异常向上抛,由 extract_content 捕获成 parse_error。"""
    from docx import Document  # 延迟导入,避免测试环境未装也能 import 本模块

    doc: DocxDocument = Document(str(file_path))
    blocks: list[TextBlock] = []
    warnings: list[str] = []
    idx = 0

    # --- 正文段落 ---
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        blocks.append(TextBlock(idx, text, "body"))
        idx += 1

    # --- 表格:每行合并为一条 ---
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if not row_text:
                continue
            blocks.append(TextBlock(idx, row_text, "table_row"))
            idx += 1

    # --- 页眉 / 页脚 (per section) ---
    for section in doc.sections:
        for para in section.header.paragraphs:
            text = para.text.strip()
            if text:
                blocks.append(TextBlock(idx, text, "header"))
                idx += 1
        for para in section.footer.paragraphs:
            text = para.text.strip()
            if text:
                blocks.append(TextBlock(idx, text, "footer"))
                idx += 1

    # --- 文本框(w:txbxContent) parser-accuracy-fixes P2-8:用预编译 lxml XPath ---
    try:
        root = doc.element.body
        txbx_nodes = _TEXTBOX_XPATH(root)
        collected: list[str] = []
        for node in txbx_nodes:
            t = (node.text or "").strip()
            if t:
                collected.append(t)
        if collected:
            combined = " ".join(collected)
            blocks.append(TextBlock(idx, combined, "textbox"))
            idx += 1
    except Exception as e:  # pragma: no cover - 文本框抽取失败只记 warning
        warnings.append(f"textbox extraction failed: {e!s}"[:200])
        logger.warning("docx textbox extract failed: %s", e)

    return DocxExtractResult(blocks=blocks, warnings=warnings)


__all__ = ["extract_docx", "DocxExtractResult", "TextBlock"]
