"""生成 C15 内置 Word 模板 default.docx

一次性脚本(开发时运行一次);结果提交到 repo。
docxtpl 占位符按 design D6 契约。

Usage:
    uv run python scripts/build_default_export_template.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


def build() -> Path:
    doc = Document()

    # Title
    t = doc.add_heading("围标检测报告", level=0)
    t.alignment = 1  # center

    # Project meta
    doc.add_paragraph("项目名称:{{ project.name }}")
    doc.add_paragraph("提交时间:{{ project.submitted_at }}")
    doc.add_paragraph("报告版本:v{{ report.version }}")

    doc.add_heading("一、综合研判", level=1)
    doc.add_paragraph("总分:{{ report.total_score }} / 100")
    doc.add_paragraph("风险等级:{{ report.risk_level }}")
    doc.add_paragraph("AI 综合研判:")
    doc.add_paragraph("{{ report.llm_conclusion }}")

    doc.add_heading("二、维度明细", level=1)
    # 采用段落级循环(比 table 行级循环更可靠)
    doc.add_paragraph("{% for dim in dimensions %}")
    doc.add_paragraph(
        "{{ dim.name }} — 最高分:{{ dim.best_score }} — "
        "铁证:{% if dim.is_ironclad %}是{% else %}否{% endif %}"
    )
    doc.add_paragraph("证据摘要:{{ dim.evidence_summary }}")
    doc.add_paragraph("{% endfor %}")

    doc.add_heading("三、高风险对比(TOP 5)", level=1)
    doc.add_paragraph("{% for pair in top_pairs %}")
    doc.add_paragraph(
        "{{ pair.bidder_a }} vs {{ pair.bidder_b }} — "
        "{{ pair.dimension }} — 分数:{{ pair.score }} — "
        "铁证:{% if pair.is_ironclad %}是{% else %}否{% endif %}"
    )
    doc.add_paragraph("摘要:{{ pair.summary }}")
    doc.add_paragraph("{% endfor %}")

    # 人工复核段:整段作为 if 块,review=None 时不输出
    doc.add_heading("四、人工复核", level=1)
    doc.add_paragraph("{% if review %}")
    doc.add_paragraph("复核结论:{{ review.status }}")
    doc.add_paragraph("评论:{{ review.comment }}")
    doc.add_paragraph(
        "复核人 ID:{{ review.reviewer_id }} — "
        "复核时间:{{ review.reviewed_at }}"
    )
    doc.add_paragraph("{% endif %}")

    out = (
        Path(__file__).parent.parent
        / "app"
        / "services"
        / "export"
        / "templates"
        / "default.docx"
    )
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


if __name__ == "__main__":
    p = build()
    print(f"wrote: {p} ({p.stat().st_size} bytes)")
