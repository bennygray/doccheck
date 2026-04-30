"""L2: tender 解析 fail-soft (detect-tender-baseline 1.16)。

覆盖 spec file-upload "tender 解析失败 fail-soft" Requirement Scenario:
- TenderDocument 解析失败 → parse_status='failed' + parse_error,**不抛异常**
- baseline_resolver 跳过 parse_status != 'extracted' 的 tender(spec 已声明,本期由
  baseline_resolver 模块自带契约约束,本测试只 verify 写入端 fail-soft)
"""

from __future__ import annotations

import io
import os
import zipfile
from pathlib import Path

from sqlalchemy import select

from app.db.session import async_session
from app.models.tender_document import TenderDocument

from ._c4_helpers import seed_project


async def test_extract_tender_archive_invalid_zip_failsoft(
    seeded_reviewer, reviewer_token, auth_client, tmp_path: Path
):
    """损坏的 zip 让 _extract_tender_archive 抛异常,但 fail-soft 写 parse_status='failed'。"""
    # 关键:本测试 INFRA_DISABLE_EXTRACT=0,触发真实异步解析
    os.environ.pop("INFRA_DISABLE_EXTRACT", None)
    project = await seed_project(owner_id=seeded_reviewer.id, name="P1")

    # 直接构造一个 TenderDocument 行,file_path 指向不存在的物理文件
    async with async_session() as s:
        bad_path = tmp_path / "nonexistent.zip"
        tender = TenderDocument(
            project_id=project.id,
            file_name="bad.zip",
            file_path=str(bad_path),
            file_size=0,
            md5="0" * 32,
            parse_status="pending",
        )
        s.add(tender)
        await s.commit()
        await s.refresh(tender)
        tender_id = tender.id

    # 直接调 _extract_tender_archive,验证 fail-soft 不抛
    from app.services.extract.engine import _extract_tender_archive

    await _extract_tender_archive(tender_id)

    # parse_status 应被设为 'failed' + parse_error 非空
    async with async_session() as s:
        row = (
            await s.execute(
                select(TenderDocument).where(TenderDocument.id == tender_id)
            )
        ).scalar_one()
        assert row.parse_status == "failed", f"got status={row.parse_status}"
        assert row.parse_error is not None
        assert "not found" in row.parse_error.lower()
        # 失败时 hash 集合应保持空(JSONB default)
        assert row.segment_hashes == []
        assert row.boq_baseline_hashes == []


async def test_extract_tender_archive_valid_zip_extracted(
    seeded_reviewer, reviewer_token, auth_client, tmp_path: Path
):
    """有效 zip 含 docx → parse_status='extracted' + segment_hashes 非空。"""
    os.environ.pop("INFRA_DISABLE_EXTRACT", None)
    project = await seed_project(owner_id=seeded_reviewer.id, name="P1")

    # 构造一个真 zip + 含 docx
    from docx import Document

    docx_buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("招标文件正文段落用于 baseline 比对的内容长度足够")
    doc.save(docx_buf)
    docx_bytes = docx_buf.getvalue()

    zip_path = tmp_path / "good.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("招标文件.docx", docx_bytes)

    async with async_session() as s:
        tender = TenderDocument(
            project_id=project.id,
            file_name="good.zip",
            file_path=str(zip_path),
            file_size=zip_path.stat().st_size,
            md5="1" * 32,
            parse_status="pending",
        )
        s.add(tender)
        await s.commit()
        await s.refresh(tender)
        tender_id = tender.id

    from app.services.extract.engine import _extract_tender_archive

    await _extract_tender_archive(tender_id)

    async with async_session() as s:
        row = (
            await s.execute(
                select(TenderDocument).where(TenderDocument.id == tender_id)
            )
        ).scalar_one()
        assert row.parse_status == "extracted", f"got status={row.parse_status}"
        assert len(row.segment_hashes) >= 1
        assert row.parse_error is None
