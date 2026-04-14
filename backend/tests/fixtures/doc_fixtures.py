"""C5 L1/L2 测试 - 在 tmp 目录构造真实的 DOCX/XLSX 样本。

C4 的 archive_fixtures 用 dummy bytes(不是合法 OOXML);C5 的 parser 需要
真实 python-docx / openpyxl 可解析的结构,所以单独一个模块。
"""

from __future__ import annotations

from pathlib import Path


def make_real_docx(
    out: Path,
    *,
    body_paragraphs: list[str] | None = None,
    header_text: str | None = None,
    footer_text: str | None = None,
    table_rows: list[list[str]] | None = None,
    author: str | None = None,
    company: str | None = None,
) -> Path:
    """用 python-docx 写一个真实的 DOCX 文件。"""
    from docx import Document

    doc = Document()
    if body_paragraphs:
        for p in body_paragraphs:
            doc.add_paragraph(p)
    if header_text:
        doc.sections[0].header.paragraphs[0].text = header_text
    if footer_text:
        doc.sections[0].footer.paragraphs[0].text = footer_text
    if table_rows:
        tbl = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r_idx, row in enumerate(table_rows):
            for c_idx, val in enumerate(row):
                tbl.cell(r_idx, c_idx).text = val
    if author:
        doc.core_properties.author = author
    # company 在 app.xml;python-docx 不直接暴露,用字典属性近似
    # 不强求,大多数场景测 author 足够

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def make_real_xlsx(
    out: Path,
    *,
    sheets: dict[str, list[list[object]]] | None = None,
    hidden_sheet: str | None = None,
) -> Path:
    """用 openpyxl 写一个真实的 XLSX,支持多 sheet 与隐藏 sheet。"""
    from openpyxl import Workbook

    wb = Workbook()
    # 删掉默认 sheet
    default = wb.active
    wb.remove(default)

    sheets = sheets or {"Sheet1": [["A", "B"], [1, 2]]}
    for name, rows in sheets.items():
        ws = wb.create_sheet(title=name)
        for row in rows:
            ws.append(row)

    if hidden_sheet and hidden_sheet in sheets:
        wb[hidden_sheet].sheet_state = "hidden"

    out.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out))
    return out


def make_price_xlsx(out: Path, row_count: int = 5) -> Path:
    """构造一个标准报价表样本 (pricing sheet)。

    sheet name = "报价清单";
    header_row = 2(行 1 标题 / 行 2 表头);
    列:A 编码 / B 名称 / C 单位 / D 数量 / E 单价 / F 合价。
    """
    rows: list[list[object]] = [
        ["某项目投标报价清单", "", "", "", "", ""],  # 行1:大标题
        ["编码", "名称", "单位", "数量", "单价", "合价"],  # 行2:表头
    ]
    for i in range(row_count):
        rows.append(
            [
                f"A{i+1:03d}",
                f"分项{i+1}",
                "m3",
                f"{(i + 1) * 10}",
                f"{1234.56 + i:.2f}",
                f"{(i + 1) * 12345:.2f}",
            ]
        )
    return make_real_xlsx(out, sheets={"报价清单": rows})


__all__ = ["make_real_docx", "make_real_xlsx", "make_price_xlsx"]
