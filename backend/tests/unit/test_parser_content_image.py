"""L1 - parser/content/image_parser 单元测试 (C5 §9.3)

DOCX 内嵌图片提取 + md5 + pHash。
"""

from __future__ import annotations

import io
from pathlib import Path

from app.services.parser.content.image_parser import extract_images_from_docx


def _build_docx_with_image(out: Path, *, embed_image: bool = True) -> Path:
    """手工构造一个含 1 张 PNG 的 DOCX。"""
    from docx import Document
    from PIL import Image

    doc = Document()
    doc.add_paragraph("含图片段落")
    if embed_image:
        # 生成一张 20x20 的红色 PNG 作为嵌入图
        img = Image.new("RGB", (20, 20), (255, 0, 0))
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        doc.add_picture(buf)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def test_no_image(tmp_path: Path) -> None:
    path = _build_docx_with_image(tmp_path / "noimg.docx", embed_image=False)
    result = extract_images_from_docx(path, tmp_path / "out")
    assert result == []


def test_single_image_md5_phash(tmp_path: Path) -> None:
    path = _build_docx_with_image(tmp_path / "img.docx", embed_image=True)
    out = tmp_path / "imgs"
    result = extract_images_from_docx(path, out)
    assert len(result) == 1
    info = result[0]
    assert len(info.md5) == 32
    assert all(c in "0123456789abcdef" for c in info.md5)
    # imagehash.phash 默认 16 字符 hex(64 bit)
    assert len(info.phash) == 16
    assert Path(info.file_path).exists()
    assert info.width == 20
    assert info.height == 20


def test_duplicate_image_dedup(tmp_path: Path) -> None:
    # 同一张图即使被 docx 引用两次,md5 相同应去重
    from docx import Document
    from PIL import Image

    doc = Document()
    img = Image.new("RGB", (10, 10), (0, 255, 0))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    for _ in range(2):
        buf.seek(0)
        doc.add_picture(buf)
    path = tmp_path / "dup.docx"
    doc.save(str(path))

    result = extract_images_from_docx(path, tmp_path / "imgs2")
    # 同 md5 只落一次
    assert len(result) == 1
